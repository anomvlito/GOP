# -*- coding: utf-8 -*-
"""
Genera la respuesta del Caso 4 (University Health Services) replicando EXACTAMENTE
el template/estilo del Caso 3 Barilla: se parte del propio .docx del Caso 3 como base
(para heredar header con logos UC, footer con numeracion, margenes y fuentes), se
cambia el texto del header y se reemplaza el cuerpo con el contenido en estilo prosa.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/home/fabian/src/gop_2026/casos"
PLANTILLA = os.path.join(BASE, "Caso 3: barilla", "Caso 3 Barilla.docx")
SALIDA_DIR = os.path.join(BASE, "Caso # 4: University Health System.", "respuesta caso")
DIAG = os.path.join(SALIDA_DIR, "diagramas")
SALIDA = os.path.join(SALIDA_DIR, "Caso 4 University.docx")

FONT = "Calibri"
SZ = Pt(10)

doc = Document(PLANTILLA)

# ---------------------------------------------------------------- HEADER
hdr = doc.sections[0].header
for p in hdr.paragraphs:
    text_runs = [r for r in p.runs if (r.text or "").strip()]
    if text_runs:
        text_runs[0].text = "Análisis de Caso 4 University Health Services"
        for r in text_runs[1:]:
            r.text = ""

# ---------------------------------------------------------------- LIMPIAR CUERPO
body = doc.element.body
sectPr = body.find(qn("w:sectPr"))
for child in list(body):
    if child.tag in (qn("w:p"), qn("w:tbl")):
        body.remove(child)

# ---------------------------------------------------------------- HELPERS
def _fmt_run(r, bold=False, italic=False, color=None):
    r.font.name = FONT
    r.font.size = SZ
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color

def para(text="", bold=False, italic=False, justify=True, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if text:
        r = p.add_run(text)
        _fmt_run(r, bold=bold, italic=italic, color=color)
    return p

def rich(parts, justify=True, space_after=6):
    """parts: lista de (texto, bold, italic)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    for tup in parts:
        txt, b, i = (tup + (False, False))[:3]
        r = p.add_run(txt)
        _fmt_run(r, bold=b, italic=i)
    return p

def subpregunta(text):
    return para(text, bold=True, justify=True, space_after=4)

def imagen(path, height_cm, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, height=Cm(height_cm))
    if caption:
        c = para(caption, italic=True, justify=False, space_after=8)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _set_borders(t):
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn("w:" + edge), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "0", qn("w:color"): "808080"})
        borders.append(el)
    tblPr.append(borders)

def tabla(headers, rows, font_sz=9):
    t = doc.add_table(rows=1, cols=len(headers))
    _set_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = ""
        pr = hdr_cells[j].paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pr.add_run(h)
        run.font.name = FONT; run.font.size = Pt(font_sz); run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = hdr_cells[j]._tc.get_or_add_tcPr()
        e = shd.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "1F4E79"})
        shd.append(e)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            pr = cells[j].paragraphs[0]
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pr.add_run(str(val))
            run.font.name = FONT; run.font.size = Pt(font_sz)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    return t

# ================================================================ CONTENIDO
rich([("Integrantes Grupo:  ", True, False),
      ("César Meneses, Fabián Ortega", False, False)], space_after=2)
para("Responda las siguientes preguntas en el espacio asignado.", space_after=10)

# ---------------------------------------------------------------- PREGUNTA 1
para("Pregunta 1 (3 puntos):", bold=True, space_after=6)

subpregunta("(1 pto.) Elabore un flujograma del sistema “antiguo” con el que atendía la "
            "clínica. ¿Dónde se generan las esperas en el sistema? ¿Cuál es el cuello de botella?")

para("En el sistema previo al triage, todo paciente seguía un recorrido estrictamente "
     "secuencial. Tras firmar la hoja numerada en recepción (2 minutos), el recepcionista "
     "solicitaba la ficha clínica al Departamento de Registros Médicos, la cual demoraba "
     "entre 8 y 9 minutos en ser trasladada físicamente hasta la clínica. Un oficinista "
     "revisaba luego la ficha durante unos 5 minutos para verificar que los reportes de "
     "laboratorio y los datos del paciente estuvieran actualizados, y la colocaba en una "
     "pila ordenada por orden de llegada. Cuando la ficha llegaba al tope de la pila, el "
     "paciente era visto por la primera enfermera disponible, quien realizaba un diagnóstico "
     "de aproximadamente 10 minutos. Si el problema era menor, lo resolvía en el acto (40% "
     "de los casos); si no, el paciente debía volver a la sala de espera para ser atendido "
     "por un médico (60% restante).")

imagen(os.path.join(DIAG, "sistema_antiguo.png"), height_cm=13.5,
       caption="Figura 1. Flujograma del sistema antiguo (pre-triage).")

para("Las esperas se generaban en tres puntos críticos. La primera y más severa ocurría "
     "entre el sign-in y el primer contacto con la enfermera, que promediaba 23 minutos, "
     "llegando a superar los 35 minutos para el 22% de los pacientes. La segunda era "
     "administrativa: la búsqueda y el traslado físico de la ficha consumían cerca de 13 "
     "minutos antes incluso de que el paciente pudiera ser evaluado. La tercera era la "
     "espera intermedia de derivación: el 60% de los pacientes que requería un médico debía "
     "volver a la cola, con un retraso adicional promedio de 10 minutos. A esto se sumaba "
     "que quienes solicitaban un médico específico (19% de la demanda) esperaban en promedio "
     "40 minutos, ya que no recibían prioridad sobre quienes habían llegado antes.")

para("El cuello de botella del sistema era la etapa de enfermería. Como el 100% de los "
     "pacientes debía pasar obligatoriamente por una consulta de 10 minutos, pero las "
     "enfermeras solo resolvían de forma definitiva el 40% de los casos, su capacidad se "
     "saturaba atendiendo evaluaciones que, en el 60% de las ocasiones, debían repetirse "
     "después en la consulta médica. Tal como señaló el Dr. Postel, cada enfermera decidía "
     "de manera independiente el alcance de la atención, lo que producía inconsistencia, "
     "alta variabilidad y, sobre todo, una duplicación de esfuerzos que congestionaba todo "
     "el flujo aguas abajo.")

subpregunta("(1 pto.) ¿Qué problema presentaba el sistema antiguo? ¿En qué métricas se "
            "reflejaba este problema? ¿Cuál es la variabilidad que tiene el sistema? "
            "Señale al menos 3 de cada una y comente por qué es un problema.")

para("El sistema antiguo presentaba al menos tres problemas de fondo. En primer lugar, "
     "tiempos de espera excesivos e impredecibles que no se correlacionaban con la urgencia "
     "ni la simplicidad de la dolencia: un paciente podía esperar casi una hora por la "
     "renovación de una receta. En segundo lugar, una duplicación ineficiente de los "
     "esfuerzos médicos, pues el 60% de los pacientes pasaba por dos evaluaciones completas "
     "e independientes (enfermera y luego médico), repitiendo preguntas y exámenes básicos. "
     "En tercer lugar, una mala percepción del servicio: la sala de espera saturada y el "
     "trato impersonal llevaban a que los pacientes describieran la clínica como “fría, "
     "ineficiente e impersonal”, e incluso a que algunos evitaran atenderse, afectando la "
     "prevención en salud.")

para("Estos problemas se reflejaban en métricas concretas: el tiempo de espera promedio "
     "al primer contacto era de 23 minutos; el 22% de los pacientes superaba los 35 minutos "
     "de espera para ver a la enfermera; la tasa de derivación a médico alcanzaba el 60%; y "
     "el tiempo de espera de derivación promediaba 10 minutos adicionales. Cada una de estas "
     "cifras evidencia un sistema cuya capacidad no lograba absorber la demanda en los "
     "períodos punta y que, además, gastaba capacidad en trabajo redundante.")

para("Respecto de la variabilidad, el sistema enfrentaba al menos tres fuentes. La primera "
     "es la variabilidad en la tasa de llegada: la demanda se concentraba en las horas "
     "intermedias del día (18,2 llegadas/hora entre 8 y 9 AM, frente a apenas 2,8 entre 5 y "
     "6 PM), lo que generaba sobrecarga matinal y ociosidad al cierre. La segunda es la "
     "variabilidad en los tiempos de servicio según la severidad: una dolencia podía "
     "resolverse en 5 minutos (un resfrío o un lavado de oído) o requerir 30 a 40 minutos "
     "(una apendicitis o un dolor de pecho). La tercera es la variabilidad por solicitudes "
     "específicas: el 19% de los pacientes exigía un proveedor en particular, lo que impedía "
     "operar con una fila única y sobrecargaba a ciertos médicos mientras otros quedaban "
     "subutilizados. Toda variabilidad, combinada con alta utilización, hace que el tiempo "
     "en cola se dispare (ecuación de Kingman), por lo que constituye el motor de las largas "
     "esperas observadas.")

subpregunta("(1 pto.) Elabore un flujograma del sistema “nuevo” con el que atendía la "
            "clínica. ¿Cuál es la mayor diferencia? ¿Cómo se hace cargo de las esperas y la "
            "variabilidad? ¿En qué métrica se debería reflejar los cambios del nuevo sistema? "
            "¿Por qué debería ser mejor?")

para("El sistema nuevo, implementado por Kathryn Angell en septiembre de 1979, introdujo "
     "un triage a la entrada. El paciente completa un formulario de visita ambulatoria (AVF) "
     "y, mientras un oficinista solicita su ficha en paralelo, es llamado por una de las dos "
     "coordinadoras de triage —enfermeras registradas altamente experimentadas— que en 3 a "
     "4 minutos determinan la naturaleza del problema y derivan al paciente al recurso "
     "adecuado: una enfermera practicante (NP) si la dolencia cae en una de las 13 "
     "categorías protocolizadas, o un médico (MD) en caso contrario o ante mayor gravedad. "
     "Los casos de emergencia se priorizan de inmediato.")

imagen(os.path.join(DIAG, "sistema_nuevo.png"), height_cm=13.0,
       caption="Figura 2. Flujograma del sistema nuevo (con triage).")

para("La mayor diferencia es la paralelización y segmentación de los flujos. Se elimina el "
     "esquema en serie (enfermera → médico) que obligaba a todos a pasar por un mismo cuello "
     "de botella, y se reemplaza por un enrutamiento que divide la demanda en dos colas "
     "paralelas según complejidad. El triage actúa como un clasificador que absorbe la "
     "variabilidad de severidad en la entrada y la canaliza al recurso correcto; además, en "
     "períodos de alta demanda las coordinadoras redirigen dinámicamente casos dudosos hacia "
     "los médicos para descongestionar a las NPs. La búsqueda física de la ficha, antes "
     "secuencial, ahora se ejecuta en paralelo mientras el paciente espera el triage, "
     "eliminando ese tiempo muerto de la ruta crítica.")

para("Los cambios deberían reflejarse en un menor tiempo total de permanencia en la clínica, "
     "en una alta tasa de resolución directa de las NPs (cercana al 95%, sin derivaciones "
     "adicionales) y en tiempos de espera más balanceados entre médicos y enfermeras "
     "practicantes. El nuevo sistema debería ser mejor porque elimina la redundancia de "
     "atención del 60% de los pacientes: al segmentar por complejidad, reserva la costosa "
     "capacidad médica para los diagnósticos avanzados y otorga autonomía a las NPs bajo "
     "guías clínicas, acortando los tiempos unitarios de procesamiento del conjunto del "
     "sistema. No obstante, conviene notar que en la práctica el beneficio fue parcial: las "
     "NPs absorbieron solo el 33% de la carga (y no más), en parte porque las coordinadoras, "
     "al ver saturadas a las practicantes, terminaban derivando más pacientes a los médicos.")

# ---------------------------------------------------------------- PREGUNTA 2
para("Pregunta 2 (3 puntos):", bold=True, space_after=6)

subpregunta("(0.5 pto) Determine la utilización del sistema desde su apertura a las 8 AM "
            "hasta las 5 PM, para el día promedio y para el lunes.")

para("La utilización se obtiene comparando la carga de trabajo demandada (en horas de "
     "consulta) con la capacidad programada (en horas-profesional) en la ventana de 8 AM a "
     "5 PM. Según el Exhibit 2, en ese tramo ingresan en promedio 140,2 pacientes al día "
     "(143 considerando el día completo), y el lunes, día punta, esa cifra escala a "
     "aproximadamente 159,8 pacientes. Bajo el sistema de triage el 67% de los pacientes "
     "termina siendo atendido por un médico y el 33% por una enfermera practicante; los "
     "médicos atienden 3,1 pacientes por hora y las NPs 1,8. La capacidad programada en el "
     "tramo se obtiene sumando las horas-profesional del Exhibit 4.")

tabla(
    ["Recurso", "Demanda (h)", "Capacidad prog. (h)", "Utilización ρ"],
    [["Médicos — día promedio", "30,3", "28,1", "108%  (inestable)"],
     ["NPs — día promedio", "25,7", "29,5", "87%"],
     ["Médicos — lunes", "34,5", "28,5", "121%  (muy inestable)"],
     ["NPs — lunes", "29,3", "29,5", "99%  (al límite)"]],
    font_sz=9)

para("El resultado es contundente: la capacidad médica está sobrepasada incluso en un día "
     "promedio (108%) y colapsa el lunes (121%), mientras que las NPs operan con holgura el "
     "día promedio (87%) pero quedan al borde de la saturación el lunes (99%). Es importante "
     "matizar que una utilización superior al 100% no significa una cola literalmente "
     "infinita: en la realidad el desbalance se absorbe porque el personal permanece después "
     "de las 6 PM despachando la cola acumulada y porque parte de la capacidad proviene de "
     "fracciones de hora y tiempos de reserva. Más relevante aún, el análisis hora a hora "
     "muestra que la capacidad está mal calzada en el tiempo: a las 8–9 AM la utilización "
     "médica supera el 190%, mientras que cerca del mediodía cae por debajo del 80%. Buena "
     "parte del problema es de timing, no solo de dotación.")

subpregunta("(0.5 pto) ¿Cuál sería la cantidad necesaria de doctores para un nivel de "
            "utilización de 100%, 90%, 75% y 60%?  (0.5 pto) ¿Cuál sería el tiempo de "
            "espera de los pacientes en cada nivel?")

para("La cantidad de médicos requerida se obtiene de la carga ofrecida, a = λ/μ. Para el "
     "día promedio, la tasa de llegada a la cola médica es λ = 93,9 pacientes / 9 horas = "
     "10,44 pac/hora, de modo que a = 10,44 / 3,1 = 3,37 médicos equivalentes para operar a "
     "100% de utilización. Para alcanzar un nivel objetivo ρ se necesita c = a/ρ médicos. "
     "El tiempo de espera en cola se estima con el modelo M/M/c (fórmula de Erlang C) "
     "evaluado en el número entero de médicos. Para el lunes, λ = 11,90 pac/hora y a = 3,84.")

tabla(
    ["Caso", "ρ objetivo", "Médicos (cont.)", "Médicos (entero)", "ρ real", "Espera Wq"],
    [["Día promedio", "100%", "3,37", "3", "112%", "∞ (inestable)"],
     ["", "90%", "3,74", "4", "84,2%", "20,6 min"],
     ["", "75%", "4,49", "5", "67,3%", "4,0 min"],
     ["", "60%", "5,61", "6", "56,1%", "1,1 min"],
     ["Lunes (punta)", "100%", "3,84", "3", "128%", "∞ (inestable)"],
     ["", "90%", "4,26", "4", "95,9%", "108,7 min"],
     ["", "75%", "5,12", "5", "76,8%", "8,2 min"],
     ["", "60%", "6,40", "6", "64,0%", "2,2 min"]],
    font_sz=9)

para("La tabla revela la enorme sensibilidad del sistema cerca de la saturación: el día "
     "promedio, pasar de 4 a 5 médicos reduce la espera de 20,6 a 4,0 minutos, y el lunes "
     "operar con 4 médicos (95,9% de utilización) dispara la espera a casi dos horas, "
     "mientras que con 5 cae a 8,2 minutos. Este comportamiento ilustra por qué nunca "
     "conviene planificar un servicio al 100% de utilización: en la vecindad de ρ = 1 el "
     "tiempo en cola crece de forma explosiva.")

subpregunta("(0.5 pto) ¿Qué haría usted? Comente la cantidad de doctores que asignaría a la "
            "clínica y por qué tomaría esa decisión.")

para("Asignaría de forma permanente 5 médicos en promedio de martes a viernes (utilización "
     "cercana al 67%, con esperas en torno a 4 minutos) y reforzaría a 6 médicos los lunes "
     "(utilización del 64%, esperas de 2 minutos). La razón es triple. Primero, evitar el "
     "colapso del lunes, donde mantener 4 médicos produce esperas de casi dos horas que son "
     "inaceptables para la meta de servicio de UHS. Segundo, disponer de un colchón de "
     "capacidad para absorber la demanda oculta: el caso reporta que las “walk-in "
     "appointments” —pacientes que sus propios médicos citan informalmente en la clínica— "
     "llegan a ocupar al 100% a dos de cada cinco médicos, reduciendo en un 40% la capacidad "
     "real disponible para los pacientes espontáneos. Operar con holgura teórica protege al "
     "sistema de esa fuga. Tercero, el costo relativo de un médico adicional (salarios de "
     "35.000 a 55.000 USD anuales más 18,5% de beneficios) es bajo frente a la mejora en "
     "satisfacción y a las horas extra que el personal de enfermería y recepción debe "
     "trabajar para despachar las colas pasadas las 6 PM.")

para("Con todo, antes de sumar dotación recomendaría reasignar las horas-médico ya "
     "existentes hacia el peak de 8 a 10 AM, donde la utilización supera el 190%, "
     "trasladándolas desde el tramo del mediodía donde la clínica está subutilizada. Calzar "
     "la oferta con la curva real de llegadas es la palanca más barata y de mayor impacto, y "
     "complementa —no reemplaza— la decisión de dotación anterior.")

subpregunta("(1 pto.) Elabore un modelo de simulación del proceso. Muestre cómo genera las "
            "esperas, las colas y las atenciones. Presente los resultados del modelo y "
            "proponga al menos dos soluciones y cuáles son los impactos en los resultados.")

para("Se construyó un modelo de simulación de eventos discretos (en Python, archivo "
     "simulacion.py) que reproduce un día de operación de la clínica. El modelo genera las "
     "llegadas mediante procesos de Poisson con las tasas horarias del Exhibit 2, hace pasar "
     "a cada paciente por la cola de triage (2 coordinadoras, 3 a 4 minutos) y luego por la "
     "cola de médicos o de NPs, con la dotación variable hora a hora del Exhibit 4 y tiempos "
     "de servicio exponenciales. Este enfoque es necesario precisamente porque, con "
     "utilización superior al 100%, el modelo analítico M/M/c entrega esperas infinitas; la "
     "simulación, en cambio, reinicia cada día y contabiliza los pacientes que quedan en "
     "cola después de las 6 PM (“rezagados”), capturando lo que realmente ocurre bajo "
     "sobrecarga transitoria. Se promediaron 3.000 días por escenario.")

tabla(
    ["Escenario", "Espera triage", "Espera MD", "Espera NP", "Rezagados/día"],
    [["Día prom. — Base", "0,53 min", "62,0 min", "24,2 min", "11,9 pac"],
     ["Día prom. — Sol. 1 (pooling)", "0,53 min", "59,8 min", "20,9 min", "11,4 pac"],
     ["Día prom. — Sol. 2 (NP 50%)", "0,53 min", "19,0 min", "100,8 min", "16,0 pac"],
     ["Lunes — Base", "0,74 min", "105,0 min", "36,5 min", "21,8 pac"],
     ["Lunes — Sol. 1 (pooling)", "0,74 min", "105,7 min", "33,0 min", "21,6 pac"],
     ["Lunes — Sol. 2 (NP 50%)", "0,74 min", "28,7 min", "152,8 min", "25,5 pac"]],
    font_sz=9)

para("La simulación confirma el diagnóstico: en el escenario base la espera para ver a un "
     "médico ronda los 62 minutos en un día promedio (y supera los 100 el lunes), muy por "
     "encima de los 24 minutos de las NPs. Esto refleja tanto el desbalance de carga como la "
     "sobrecarga estructural del recurso médico (ρ ≈ 108%), y explica los casi 12 pacientes "
     "diarios que siguen en el sistema al cierre. A partir de aquí se evaluaron dos "
     "soluciones. La primera, eliminar la posibilidad de solicitar un médico específico "
     "(pooling perfecto de la capacidad médica), reduce la espera de los médicos de 62 a 60 "
     "minutos sin costo alguno de contratación; el efecto del pooling es real pero modesto, "
     "coherente con que el problema de fondo es la sobrecarga agregada y no solo la ausencia "
     "de una fila única. La segunda, expandir las guías clínicas de las NPs para que absorban "
     "el 50% de los pacientes en lugar del 33%, reduce drásticamente la espera médica (a 19 "
     "minutos), pero —si no se aumenta la dotación de NPs— traslada el cuello de botella "
     "hacia ellas y hace explotar su cola a más de 100 minutos, con cerca de 16 rezagados "
     "diarios. La conclusión es que esta "
     "segunda medida solo es viable si se acompaña de un incremento equivalente de la "
     "capacidad de las NPs (contratando o liberando sus horas administrativas); de lo "
     "contrario, simplemente se traslada el problema en vez de resolverlo. La recomendación "
     "final combina ambas palancas con el recalce horario de la dotación médica hacia el "
     "peak matinal.")

# ---------------------------------------------------------------- sectPr al final
if sectPr is not None:
    body.remove(sectPr)
    body.append(sectPr)

doc.save(SALIDA)
print("OK ->", SALIDA)
print("Párrafos en cuerpo:", len(doc.paragraphs), "| Tablas:", len(doc.tables))
